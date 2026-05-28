import io
import docx
from pypdf import PdfReader
from app.core.exceptions import ExtractionException

class ResumeExtractor:
    """Service to handle extracting raw text from PDF and DOCX resume formats."""

    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """Extracts text from a PDF file using pypdf.
        
        Preserves simple layouts and page boundaries.
        """
        try:
            pdf_file = io.BytesIO(file_bytes)
            reader = PdfReader(pdf_file)
            
            extracted_pages = []
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    # Clean up trailing spaces per line to preserve paragraph formatting
                    cleaned_lines = [line.rstrip() for line in text.split("\n")]
                    page_text = "\n".join(cleaned_lines)
                    extracted_pages.append(f"--- [Page {page_num + 1}] ---\n{page_text}")
            
            if not extracted_pages:
                raise ExtractionException("PDF is empty or could not be read (scanned/image-only PDFs not supported without OCR).")
            
            return "\n\n".join(extracted_pages)
        except Exception as e:
            if isinstance(e, ExtractionException):
                raise e
            raise ExtractionException(f"Failed to parse PDF document: {str(e)}")

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """Extracts text from a DOCX file using python-docx."""
        try:
            docx_file = io.BytesIO(file_bytes)
            doc = docx.Document(docx_file)
            
            paragraphs = []
            # Extract standard paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # Extract text from tables if any exist (extremely common in resume templates)
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        # Join cells with vertical separators to capture tabular metadata
                        table_text.append(" | ".join(dict.fromkeys(row_data)))
            
            full_text = "\n".join(paragraphs)
            if table_text:
                full_text += "\n\n--- Document Tables ---\n" + "\n".join(table_text)
                
            if not full_text.strip():
                raise ExtractionException("DOCX file contains no readable text content.")
                
            return full_text
        except Exception as e:
            if isinstance(e, ExtractionException):
                raise e
            raise ExtractionException(f"Failed to parse DOCX document: {str(e)}")

    def extract_text(self, file_bytes: bytes, filename: str) -> str:
        """Helper to dynamically route extraction based on file extension."""
        ext = filename.split(".")[-1].lower()
        if ext == "pdf":
            return self.extract_text_from_pdf(file_bytes)
        elif ext in ["docx", "doc"]:
            return self.extract_text_from_docx(file_bytes)
        else:
            raise ExtractionException(f"Unsupported file format: '.{ext}'. Only PDF and DOCX are supported.")
